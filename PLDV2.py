import base64
import datetime
import io
import os
import re
import zipfile
from docx import Document
import pandas as pd
from PIL import Image
import streamlit as st

# 1. Configuração da página
st.set_page_config(
    page_title="Sistema de Dossiês PLD-FT",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# Otimização da imagem de fundo
@st.cache_data
def get_optimized_base64_background(bin_file):
    if os.path.exists(bin_file):
        try:
            img = Image.open(bin_file)
            img.thumbnail((1920, 1080))
            buffer = io.BytesIO()
            if img.mode != "RGB":
                img = img.convert("RGB")
            img.save(buffer, format="JPEG", quality=75, optimize=True)
            return base64.b64encode(buffer.getvalue()).decode()
        except Exception:
            with open(bin_file, "rb") as f:
                return base64.b64encode(f.read()).decode()
    return None


NOME_IMAGEM_FUNDO = (
    "abstract-metallic-wave-texture-with-glossy-reflective-surface-dark-lighting.jpg"
)
bin_str = get_optimized_base64_background(NOME_IMAGEM_FUNDO)

bg_css = (
    f"""
    .stApp {{
        background: linear-gradient(rgba(9, 9, 11, 0.80), rgba(9, 9, 11, 0.80)),
                    url("data:image/jpeg;base64,{bin_str}") no-repeat center center fixed !important;
        background-size: cover !important;
    }}
"""
    if bin_str
    else ".stApp { background-color: #09090b !important; }"
)

# 2. CSS Estilizado
st.markdown(
    f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Bricolage+Grotesque:opsz,wght@12..96,200..800&family=Inter:wght@400;600;700&display=swap');

    {bg_css}

    .stApp h1 {{
        font-family: 'Bricolage Grotesque', sans-serif !important;
        font-weight: 800 !important;
        font-size: 2.5rem !important;
        color: #FFFFFF !important;
    }}

    .stApp h2, .stApp h3, .stApp h4, .stApp label, .stApp .stMarkdown p {{
        font-family: 'Bricolage Grotesque', sans-serif !important;
        color: #F4F4F5 !important;
    }}

    [data-testid="stFileUploader"] > div {{
        background-color: rgba(24, 24, 27, 0.85) !important;
        border: 1px dashed rgba(255, 255, 255, 0.3) !important;
        border-radius: 12px !important;
        padding: 1.5rem !important;
    }}

    .stTabs [data-baseweb="tab-list"] {{
        gap: 8px;
        background-color: rgba(18, 18, 18, 0.75);
        padding: 6px;
        border-radius: 10px;
        border: 1px solid rgba(255, 255, 255, 0.15);
    }}

    .stTabs [data-baseweb="tab"] {{
        height: 44px;
        background-color: transparent;
        border-radius: 6px;
        color: #A1A1AA !important;
        font-family: 'Bricolage Grotesque', sans-serif !important;
        font-weight: 600;
        font-size: 1rem;
        padding: 0 16px;
        border: none !important;
    }}

    .stTabs [aria-selected="true"] {{
        background-color: #18181B !important;
        color: #FFFFFF !important;
        border: 1px solid rgba(255, 255, 255, 0.3) !important;
    }}

    .stTextInput input, .stSelectbox select, .stTextArea textarea {{
        background-color: rgba(15, 15, 18, 0.9) !important;
        color: #FFFFFF !important;
        border: 1px solid rgba(255, 255, 255, 0.2) !important;
        border-radius: 8px !important;
        font-family: 'Inter', sans-serif !important;
    }}

    .stButton > button {{
        width: 100%;
        background: linear-gradient(180deg, #27272A 0%, #09090B 100%);
        color: #FFFFFF !important;
        font-family: 'Bricolage Grotesque', sans-serif !important;
        font-weight: 700;
        border-radius: 8px;
        border: 1px solid rgba(255, 255, 255, 0.25);
    }}
    </style>
    """,
    unsafe_allow_html=True,
)


def gerar_codigo_dossie(indice):
    hoje = datetime.date.today().strftime("%Y%m%d")
    return f"DOS-{hoje}-{str(indice).zfill(3)}"


def formatar_data(valor):
    if pd.isna(valor) or not valor:
        return ""
    val_str = str(valor).strip()
    match = re.search(r"(\d{4})-(\d{2})-(\d{2})", val_str)
    if match:
        ano, mes, dia = match.groups()
        return f"{dia}/{mes}/{ano}"
    return val_str


def formatar_moeda(valor):
    if pd.isna(valor) or not valor:
        return "R$ 0,00"
    try:
        val_float = float(str(valor).replace(",", "."))
        return (
            f"R$ {val_float:,.2f}"
            .replace(",", "v")
            .replace(".", ",")
            .replace("v", ".")
        )
    except Exception:
        return f"R$ {valor}"


def preencher_tabela_diligencias(doc_obj, lista_diligencias, datas_diligencias):
    """Insere as linhas na tabela de diligências e remove a linha template."""
    for table in doc_obj.tables:
        for row in table.rows:
            cell_texts = [c.text for c in row.cells]
            if any("{{DILIGENCIAS_NOME}}" in t for t in cell_texts):
                for dil in lista_diligencias:
                    dt = datas_diligencias.get(
                        dil, datetime.date.today().strftime("%d/%m/%Y")
                    )
                    new_row = table.add_row()
                    if len(new_row.cells) >= 2:
                        new_row.cells[0].text = str(dil)
                        new_row.cells[1].text = str(dt)
                    elif len(new_row.cells) == 1:
                        new_row.cells[0].text = f"{dil} - {dt}"

                tr = row._tr
                table._tbl.remove(tr)
                break


def substituir_texto(doc_obj, mapa_substituicao):
    """Substitui as variáveis {{TAG}} no corpo e nas tabelas do documento."""
    for p in doc_obj.paragraphs:
        for chave, valor in mapa_substituicao.items():
            if chave in p.text:
                for run in p.runs:
                    if chave in run.text:
                        run.text = run.text.replace(chave, str(valor))
                if chave in p.text:
                    p.text = p.text.replace(chave, str(valor))

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for chave, valor in mapa_substituicao.items():
                    if chave in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if chave in run.text:
                                    run.text = run.text.replace(
                                        chave, str(valor)
                                    )
                            if chave in p.text:
                                p.text = p.text.replace(chave, str(valor))


# --- CABEÇALHO DA PLATAFORMA ---
col_logo, col_titulo = st.columns([1.2, 5], vertical_alignment="center")
with col_logo:
    if os.path.exists("noBgWhite.png"):
        st.image("noBgWhite.png", width=180)
    else:
        st.write("🛡️")

with col_titulo:
    st.title("Sistema de Gestão e Emissão de Dossiês PLD-FT")
    st.caption(
        "Plataforma de Automação de Análise de Alertas, Diligências e Compliance"
    )

st.markdown("---")

tab1, tab2, tab3 = st.tabs(
    [
        "📂 1. Carregamento & Alertas",
        "🔍 2. Análise & Diligências",
        "📄 3. Emissão do Dossiê",
    ]
)

# --- ABA 1: UPLOAD & SELEÇÃO ---
with tab1:
    st.markdown("### 1. Upload da Planilha de Ocorrências")
    uploaded_file = st.file_uploader(
        "Arraste e solte a base de detectados (.xlsx ou .csv):",
        type=["xlsx", "xls", "csv"],
    )

    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith(".csv"):
                df = pd.read_csv(uploaded_file, dtype=str)
            else:
                df = pd.read_excel(uploaded_file, dtype=str)

            df.columns = [str(c).strip() for c in df.columns]
            df = df.fillna("")

            df["CODIGO_DOSSIE"] = [
                gerar_codigo_dossie(i + 1) for i in range(len(df))
            ]

            col_cpf = (
                "CPF/CNPJ Pesquisado"
                if "CPF/CNPJ Pesquisado" in df.columns
                else df.columns[6]
            )
            col_nome = (
                "Nome Encontrado"
                if "Nome Encontrado" in df.columns
                else df.columns[8]
            )

            df["ID_Alerta"] = df.apply(
                lambda r: f"{r['CODIGO_DOSSIE']} | {r.get(col_nome, '')} (CPF/CNPJ: {r.get(col_cpf, '')})",
                axis=1,
            )

            st.session_state["df_pld"] = df
            st.success(
                f"✅ Base carregada com sucesso! **{len(df)} alerta(s)** prontos para análise."
            )

            st.markdown("---")
            st.markdown("### Seleção do Registro para Análise")
            alerta_selecionado = st.selectbox(
                "Escolha o alerta para trabalhar nesta sessão:",
                df["ID_Alerta"].tolist(),
            )
            st.session_state["alerta_selecionado"] = alerta_selecionado

        except Exception as e:
            st.error(f"Erro ao ler/processar a planilha: {e}")

# Processamento Principal
if "df_pld" in st.session_state and "alerta_selecionado" in st.session_state:
    df = st.session_state["df_pld"]
    alerta_selecionado = st.session_state["alerta_selecionado"]
    linha = df[df["ID_Alerta"] == alerta_selecionado].iloc[0]

    # Identificação flexível de colunas
    col_dt_hit = next(
        (
            c
            for c in df.columns
            if "detec" in c.lower() or "hit" in c.lower() or "geraç" in c.lower()
        ),
        "Data da Detecção do Hit",
    )
    col_cpf = next(
        (
            c
            for c in df.columns
            if "cpf" in c.lower() or "cnpj" in c.lower() or "pesquisado" in c.lower()
        ),
        "CPF/CNPJ Pesquisado",
    )
    col_nome = next(
        (
            c
            for c in df.columns
            if "nome" in c.lower() or "encontrado" in c.lower()
        ),
        "Nome Encontrado",
    )

    op_origem = linha.get("Nome do Cliente", "")
    op_data = formatar_data(linha.get("Data da Operação", ""))
    op_valor = formatar_moeda(linha.get("Valor da Operação", ""))
    data_geracao = formatar_data(linha.get(col_dt_hit, ""))
    cpf_cnpj = linha.get(col_cpf, "")
    status_ip = linha.get("Parte Relacionada", "")
    nome_contraparte = linha.get(col_nome, "")
    regra_lista = linha.get("Lista", "")
    obs_complemento = linha.get("Complemento", "")
    op_destino = nome_contraparte

    modelos_justificativas = {
        "Arquivado - Sem Indício de Irregularidade": f"Análise realizada sobre o apontamento na lista '{regra_lista}'. Consultas efetuadas nas fontes abertas e bases públicas não identificaram risco iminente de PLD-FT ou atipicidade financeira.",
        "Arquivado - Falso Positivo / Homônimo": f"Análise efetuada indica tratar-se de Falso Positivo / Homônimo. Os dados cadastrais do pesquisado não coincidem com o indivíduo/entidade registrado na lista restritiva '{regra_lista}'.",
        "Encaminhado para Comunicação (COAF)": f"Constatados indícios de atipicidade e divergência cadastral incompatíveis com a capacidade financeira. Processo encaminhado ao Comitê para deliberação de Comunicação de Atipicidade ao COAF.",
        "Em Monitoramento Contínuo": f"O cliente/contraparte permanece sob monitoramento contínuo reforçado para verificação de novas transações e eventual evolução dos apontamentos na lista '{regra_lista}'.",
        "Outro (Especifique na Justificativa)": "",
    }

    # --- ABA 2: ANÁLISE E DILIGÊNCIAS ---
    with tab2:
        c1, c2 = st.columns(2)

        with c1:
            st.markdown("#### 📌 Dados do Alerta (Base Detectada)")
            st.text_input(
                "Nº Alerta / Rastreabilidade",
                linha.get("CODIGO_DOSSIE"),
                disabled=True,
            )
            st.text_input(
                "Contraparte / Destino", nome_contraparte, disabled=True
            )
            st.text_input("CPF/CNPJ Pesquisado", cpf_cnpj, disabled=True)
            st.text_input(
                "Regra / Lista Restritiva", regra_lista, disabled=True
            )
            st.text_input("Status na IP", status_ip, disabled=True)
            st.text_input("Operação - Origem", op_origem, disabled=True)
            st.text_input("Operação - Data", op_data, disabled=True)
            st.text_input("Operação - Valor", op_valor, disabled=True)

        with c2:
            st.markdown("#### ⚖️ Decisão & Matriz de Risco")
            analista = st.text_input(
                "Analista Responsável",
                "Analista PLD",
                key="input_analista",
            )
            data_analise = st.date_input(
                "Data da Análise", datetime.date.today(), key="input_data_an"
            ).strftime("%d/%m/%Y")

            risco_cliente = st.selectbox(
                "Classificação de Risco do Cliente:",
                ["Baixo", "Médio", "Alto", "Não Classificado"],
                index=0,
                key="select_risco",
            )

            decisao_arquivamento = st.selectbox(
                "Conclusão da Análise:",
                list(modelos_justificativas.keys()),
                key="select_decisao",
            )

            st.markdown("---")
            st.markdown("##### 🔎 Diligências Efetuadas")

            diligencias_opcoes = st.multiselect(
                "Selecione as diligências padrão:",
                [
                    "Consulta Mídia Negativa",
                    "Pesquisa de Bens / Cartório",
                    "Consulta Base Pública (Receita / Sanções / CEIS)",
                    "Solicitação de Esclarecimentos ao Cliente",
                    "Verificação de Vínculos / Relacionamento",
                ],
                default=["Consulta Base Pública (Receita / Sanções / CEIS)"],
                key="select_diligencias",
            )

            diligencias_extras_raw = st.text_area(
                "➕ Adicionar diligências personalizadas (uma por linha):",
                placeholder="Análise de Contrato Social na Junta Comercial\nConsulta ao Sintegra / Cadastro Estadual\nVerificação de PEP no Portal da Transparência",
                height=100,
                key="input_dil_extra",
            )

            lista_final_diligencias = list(diligencias_opcoes)
            if diligencias_extras_raw.strip():
                extras = [
                    l.strip()
                    for l in diligencias_extras_raw.splitlines()
                    if l.strip()
                ]
                lista_final_diligencias.extend(extras)

            datas_diligencias = {}
            if lista_final_diligencias:
                st.markdown("**Datas da realização:**")
                for dil in lista_final_diligencias:
                    d_data = st.date_input(
                        f"Data - {dil}:",
                        datetime.date.today(),
                        key=f"data_{dil}",
                    )
                    datas_diligencias[dil] = d_data.strftime("%d/%m/%Y")

            st.markdown("---")
            texto_padrao = modelos_justificativas.get(decisao_arquivamento, "")

            justificativa = st.text_area(
                "Justificativa da Decisão (Editável):",
                value=texto_padrao,
                height=120,
                key="input_justificativa",
            )

    # --- ABA 3: EMISSÃO DO DOSSIÊ ---
    with tab3:
        st.markdown("### 📄 Emissão e Exportação do Relatório Oficial")

        col_ind, col_lote = st.columns(2, gap="large")

        # Data por extenso
        meses = [
            "janeiro",
            "fevereiro",
            "março",
            "abril",
            "maio",
            "junho",
            "julho",
            "agosto",
            "setembro",
            "outubro",
            "novembro",
            "dezembro",
        ]
        hoje = datetime.date.today()
        data_hoje_extenso = (
            f"São Paulo, {hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"
        )

        # OPÇÃO 1: DOWNLOAD INDIVIDUAL
        with col_ind:
            st.markdown("#### 👤 Download do Dossiê Selecionado")
            st.info(
                f"**Alerta:** {linha.get('CODIGO_DOSSIE')}\n\n**Contraparte:** {nome_contraparte}"
            )

            if st.button("🚀 Gerar Dossiê do Alerta Atual"):
                if os.path.exists("modelo_dossie.docx"):
                    doc = Document("modelo_dossie.docx")
                else:
                    doc = Document()

                if lista_final_diligencias:
                    preencher_tabela_diligencias(
                        doc, lista_final_diligencias, datas_diligencias
                    )

                dicionario_dados = {
                    "{{CODIGO_DOSSIE}}": linha.get("CODIGO_DOSSIE", ""),
                    "{{NUM_ALERTA}}": linha.get("CODIGO_DOSSIE", ""),
                    "{{SISTEMA}}": "Advice e-Guardian",
                    "{{NORMATIVA}}": (
                        "Lei nº 9.613/1998 e Resolução BCB nº 96/2021"
                    ),
                    "{{DATA_GERACAO}}": data_geracao,
                    "{{DATA_ELABORACAO}}": data_hoje_extenso,
                    "{{CPF_CNPJ}}": cpf_cnpj,
                    "{{NOME_CONTRAPARTE}}": nome_contraparte,
                    "{{REGRA}}": regra_lista,
                    "{{TIPOLOGIA}}": regra_lista,
                    "{{STATUS_IP}}": status_ip,
                    "{{OBS_CONTRAPARTE}}": obs_complemento,
                    "{{OPERAÇÃO_ORIGEM}}": op_origem,
                    "{{OPERAÇÃO_DESTINO}}": op_destino,
                    "{{OPERAÇÃO_DATA}}": op_data,
                    "{{OPERAÇÃO_VALOR}}": op_valor,
                    "{{RISCO_CLIENTE}}": risco_cliente,
                    "{{ANALISTA}}": analista,
                    "{{DATA_ANALISE}}": data_analise,
                    "{{STATUS_ALERTA}}": decisao_arquivamento,
                    "{{DECISAO}}": decisao_arquivamento,
                    "{{JUSTIFICATIVA}}": justificativa,
                }

                substituir_texto(doc, dicionario_dados)

                cod_dossie = linha.get("CODIGO_DOSSIE", "DOSSIE")
                nome_arquivo = f"Dossiê PLD - {cod_dossie}.docx"

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label=f"📥 Baixar Dossiê (.docx)",
                    data=buffer,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        # OPÇÃO 2: DOWNLOAD EM LOTE (.ZIP) - CORRIGIDO
        with col_lote:
            st.markdown("#### 📦 Download em Lote (Todos os Alertas)")
            st.warning(
                f"Serão gerados **{len(df)} dossiês** em um arquivo compactado (.ZIP) aplicando as configurações e diligências da análise."
            )

            if st.button("⚡ Gerar Pacote em Lote (.ZIP)"):
                with st.spinner("Gerando todos os dossiês em lote..."):
                    zip_buffer = io.BytesIO()

                    with zipfile.ZipFile(
                        zip_buffer, "w", zipfile.ZIP_DEFLATED
                    ) as zip_file:
                        for idx, row in df.iterrows():
                            if os.path.exists("modelo_dossie.docx"):
                                doc_item = Document("modelo_dossie.docx")
                            else:
                                doc_item = Document()

                            # 1. Preenche tabela de diligências no documento do lote
                            if lista_final_diligencias:
                                preencher_tabela_diligencias(
                                    doc_item,
                                    lista_final_diligencias,
                                    datas_diligencias,
                                )

                            cod = row.get("CODIGO_DOSSIE", f"DOS-{idx+1}")
                            item_cpf = row.get(col_cpf, "")
                            item_nome = row.get(col_nome, "")
                            item_regra = row.get("Lista", "")
                            item_dt_ger = formatar_data(row.get(col_dt_hit, ""))
                            item_status = row.get("Parte Relacionada", "")
                            item_obs = row.get("Complemento", "")
                            item_op_origem = row.get("Nome do Cliente", "")
                            item_op_data = formatar_data(
                                row.get("Data da Operação", "")
                            )
                            item_op_val = formatar_moeda(
                                row.get("Valor da Operação", "")
                            )

                            # Se a justificativa padrão mudar por conta da regra específica do item do lote
                            justificativa_item = justificativa
                            if (
                                justificativa
                                == modelos_justificativas.get(
                                    decisao_arquivamento, ""
                                )
                            ):
                                justificativa_item = (
                                    f"Análise realizada sobre o apontamento na lista '{item_regra}'. Consultas efetuadas nas fontes abertas e bases públicas não identificaram risco iminente de PLD-FT ou atipicidade financeira."
                                    if decisao_arquivamento
                                    == "Arquivado - Sem Indício de Irregularidade"
                                    else justificativa
                                )

                            dic_item = {
                                "{{CODIGO_DOSSIE}}": cod,
                                "{{NUM_ALERTA}}": cod,
                                "{{SISTEMA}}": "Advice e-Guardian",
                                "{{NORMATIVA}}": (
                                    "Lei nº 9.613/1998 e Resolução BCB nº"
                                    " 96/2021"
                                ),
                                "{{DATA_GERACAO}}": item_dt_ger,
                                "{{DATA_ELABORACAO}}": data_hoje_extenso,
                                "{{CPF_CNPJ}}": item_cpf,
                                "{{NOME_CONTRAPARTE}}": item_nome,
                                "{{REGRA}}": item_regra,
                                "{{TIPOLOGIA}}": item_regra,
                                "{{STATUS_IP}}": item_status,
                                "{{OBS_CONTRAPARTE}}": item_obs,
                                "{{OPERAÇÃO_ORIGEM}}": item_op_origem,
                                "{{OPERAÇÃO_DESTINO}}": item_nome,
                                "{{OPERAÇÃO_DATA}}": item_op_data,
                                "{{OPERAÇÃO_VALOR}}": item_op_val,
                                "{{RISCO_CLIENTE}}": risco_cliente,
                                "{{ANALISTA}}": analista,
                                "{{DATA_ANALISE}}": data_analise,
                                "{{STATUS_ALERTA}}": decisao_arquivamento,
                                "{{DECISAO}}": decisao_arquivamento,
                                "{{JUSTIFICATIVA}}": justificativa_item,
                            }

                            substituir_texto(doc_item, dic_item)

                            doc_buf = io.BytesIO()
                            doc_item.save(doc_buf)
                            doc_buf.seek(0)

                            nome_limpo_contraparte = re.sub(
                                r'[\\/*?:"<>|]', "", item_nome[:25].strip()
                            )
                            fname = f"Dossiê PLD - {cod} - {nome_limpo_contraparte}.docx"
                            zip_file.writestr(fname, doc_buf.getvalue())

                    zip_buffer.seek(0)
                    nome_zip = f"Dossies_PLD_{datetime.date.today().strftime('%Y%m%d')}.zip"

                    st.download_button(
                        label=f"📥 Baixar Arquivo .ZIP ({len(df)} Dossiês)",
                        data=zip_buffer,
                        file_name=nome_zip,
                        mime="application/zip",
                    )
else:
    with tab2:
        st.warning(
            "⚠️ Faça o upload da planilha na **Aba 1** para habilitar a análise."
        )
    with tab3:
        st.warning("⚠️ Nenhum alerta carregado para emissão.")
