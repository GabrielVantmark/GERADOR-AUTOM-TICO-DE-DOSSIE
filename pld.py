import datetime
from docx import Document
import streamlit as st
from PIL import Image

# Carrega a imagem da logo
favicon = Image.open("favicon.png")

# 2. Injeção de CSS para expandir a largura da página e diminuir as margens
st.markdown(
    """
    <style>
    /* Remove as margens excessivas do topo e das laterais */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        padding-left: 3rem;
        padding-right: 3rem;
        max-width: 95% !important; /* Expande para ocupar 95% da tela */
    }
    </style>
    """,
    unsafe_allow_html=True  # <--- Aqui entra diretamente a palavra True com 'T' maiúsculo
)

# --- CABEÇALHO COM LOGO E TÍTULO ALINHADOS ---
col_logo, col_titulo = st.columns([1, 6], vertical_alignment="center")  # Proporção 1:6 para a logo não ficar gigante

with col_logo:
    # Substitua "logo.png" pelo nome/caminho do arquivo da sua logo
    try:
        st.image("logo title.png", width=150)  # Ajuste a largura (width) conforme necessário
    except Exception:
        st.write("🛡️")  # Fallback caso a imagem não seja encontrada

with col_titulo:
    st.title("Gerador Automático de Dossiê PLD-FT")
    st.subheader("Otimize a documentação de análises do e-Guardian")

st.markdown("---")

st.set_page_config(page_icon=favicon)


# --- FORMULÁRIO DE ENTRADA DE DADOS ---
st.markdown("### 1. Informações do Alerta")
col1, col2, col3 = st.columns(3)

with col1:
    num_alerta = st.text_input("Nº do Alerta", value="", placeholder="Ex: 38499387861")
    analista = st.text_input("Analista Responsável", value="", placeholder="Nome do analista")

with col2:
    data_geracao = st.date_input(
        "Data de Geração do Alerta", 
        datetime.date.today(),
        format="DD/MM/YYYY"
    )
    status_alerta = st.selectbox(
        "Status do Alerta",
        [
            "Analisado/arquivado positivo",
            "Analisado/arquivado falso positivo",
            "Em análise",
        ],
    )

with col3:
    data_analise = st.date_input(
        "Data da Análise", 
        datetime.date.today(),
        format="DD/MM/YYYY"
    )
    sistema_monitoramento = st.text_input("Sistema de Monitoramento", value="Advice e-Guardian")

st.markdown("---")
st.markdown("### 2. Tipologia e Contraparte")
col4, col5 = st.columns(2)

with col4:
    tipologia = st.text_input(
        "Tipologia Identificada",
        value="",
        placeholder="Ex: Contraparte constante em lista restritiva",
    )
    regra_disparada = st.text_input(
        "Regra de Monitoramento Disparada",
        value="",
        placeholder="Ex: Presença em lista restritiva",
    )
    normativa = st.text_input(
        "Normativa de Referência",
        value="",
        placeholder="Ex: Lei nº 9.613/1998 e Resolução BCB nº 96/2021",
    )

with col5:
    nome_contraparte = st.text_input("Nome / Razão Social", value="", placeholder="Nome completo da pessoa ou empresa")
    cpf_cnpj = st.text_input("CPF / CNPJ", value="", placeholder="000.000.000-00")
    obs_contraparte = st.text_area(
        "Observação sobre a Contraparte (Envolvimento)",
        value="",
        placeholder="Detalhes sobre o envolvimento da contraparte...",
        height=100,
    )

st.markdown("---")
st.markdown("### 3. Diligências Realizadas e Datas")

opcoes_diligencias = [
    "Análise do histórico transacional da contraparte na base de dados da MIUPAG.",
    "Registro da análise do caso em dossiê formalizado e ciência para alta Administração.",
    "Consulta as bases públicas e privadas para devida verificação da correlação entre a contraparte e a empresa citada em investigação.",
]

diligencias_selecionadas = st.multiselect(
    "Selecione as diligências executadas:",
    options=opcoes_diligencias,
    default=opcoes_diligencias,
)

diligencia_personalizada = st.text_area(
    "Outras Diligências (opcional / digite para adicionar):",
    value="",
    placeholder="Digite uma diligência adicional aqui...",
    height=80,
)

# Unifica todas as diligências escolhidas/digitadas
lista_todas_diligencias = list(diligencias_selecionadas)
if diligencia_personalizada.strip():
    linhas_manuais = [l.strip() for l in diligencia_personalizada.split("\n") if l.strip()]
    lista_todas_diligencias.extend(linhas_manuais)

# Dicionário que guardará o par: (Texto da Diligência, Data Formatada DD/MM/AAAA)
diligencias_com_datas = []

if lista_todas_diligencias:
    st.write("**Ajuste as datas de execução para cada diligência:**")
    for idx, item in enumerate(lista_todas_diligencias):
        col_text, col_date = st.columns([3, 1])
        with col_text:
            st.caption(f"📍 {item}")
        with col_date:
            data_item = st.date_input(
                f"Data da diligência #{idx+1}",
                value=data_analise,
                format="DD/MM/YYYY",
                key=f"data_dil_{idx}",
                label_visibility="collapsed"
            )
        diligencias_com_datas.append((item, data_item.strftime("%d/%m/%Y")))

st.markdown("---")
st.markdown("### 4. Decisão e Comunicação COAF")
col6, col7 = st.columns(2)

with col6:
    tipo_arquivamento = st.radio(
        "Tipo de Arquivamento", ["Positivo", "Negativo (Falso Positivo)"]
    )
    comunicado_coaf = st.radio("Comunicado ao COAF?", ["Não", "Sim", "Em avaliação"])

with col7:
    justificativa = st.text_area(
        "Justificativa do Arquivamento",
        value="",
        placeholder="Digite a justificativa da análise...",
        height=120,
    )

# --- BOTÃO DE GERAÇÃO DO DOSSIÊ ---
if st.button("🚀 Gerar Dossiê Word", type="primary"):
    if not num_alerta:
        st.warning("⚠️ Por favor, preencha pelo menos o número do alerta antes de gerar.")
    else:
        try:
            doc = Document("modelo_dossie.docx")

            # Preenche a tabela de diligências com suas respectivas datas (DD/MM/AAAA)
            for table in doc.tables:
                if len(table.rows) > 0 and "Diligência Executada" in table.rows[0].cells[0].text:
                    while len(table.rows) > 1:
                        tr = table.rows[1]._tr
                        table._tbl.remove(tr)
                    
                    for item_texto, item_data in diligencias_com_datas:
                        nova_linha = table.add_row()
                        nova_linha.cells[0].text = item_texto
                        nova_linha.cells[1].text = item_data

            # Substituição dos demais placeholders
            def substituir_texto(doc_obj, mapa_substituicao):
                for p in doc_obj.paragraphs:
                    for chave, valor in mapa_substituicao.items():
                        if chave in p.text:
                            p.text = p.text.replace(chave, str(valor))

                for table in doc_obj.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for chave, valor in mapa_substituicao.items():
                                if chave in cell.text:
                                    cell.text = cell.text.replace(chave, str(valor))

            dicionario_dados = {
                "{{NUM_ALERTA}}": num_alerta,
                "{{DATA_GERACAO}}": data_geracao.strftime("%d/%m/%Y"),
                "{{ANALISTA}}": analista,
                "{{DATA_ANALISE}}": data_analise.strftime("%d/%m/%Y"),
                "{{SISTEMA}}": sistema_monitoramento,
                "{{STATUS_ALERTA}}": status_alerta,
                "{{TIPOLOGIA}}": tipologia,
                "{{REGRA}}": regra_disparada,
                "{{NORMATIVA}}": normativa,
                "{{NOME_CONTRAPARTE}}": nome_contraparte,
                "{{CPF_CNPJ}}": cpf_cnpj,
                "{{OBS_CONTRAPARTE}}": obs_contraparte,
                "{{JUSTIFICATIVA}}": justificativa,
            }

            substituir_texto(doc, dicionario_dados)

            nome_arquivo = f"Dossie_PLD_{num_alerta}.docx"
            doc.save(nome_arquivo)

            with open(nome_arquivo, "rb") as file:
                st.download_button(
                    label="📥 Baixar Dossiê Gerado",
                    data=file,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            st.success("Dossiê gerado com sucesso!")

        except FileNotFoundError:
            st.error(
                "Arquivo 'modelo_dossie.docx' não encontrado na pasta do projeto."
            )